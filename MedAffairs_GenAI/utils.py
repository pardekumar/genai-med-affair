from langchain.vectorstores import FAISS
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
import os
from PyPDF2 import PdfReader
from docx import Document
from langchain.document_loaders import PyPDFLoader
from langchain.document_loaders import TextLoader
from config import OPENAI_API_KEY, chain_type
from langchain.embeddings import OpenAIEmbeddings
import numpy as np
from nltk.translate.bleu_score import corpus_bleu
from rouge_score import rouge_scorer
import openai
import pdfplumber

embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)


def create_page_break(document):
    """Adds page break in a document

    :param document: document in which a page break should be added
    """
    document.add_page_break()


def iterate_document_sections(document):
    """Iterate through the paragraphs in a document

    :param document: document in which paragraphs should be yielded
    """
    paragraphs = [document.paragraphs[0]]
    for paragraph in document.paragraphs[1:]:
        if is_heading(paragraph):
            yield paragraphs
            paragraphs = [paragraph]
            continue
        paragraphs.append(paragraph)
    yield paragraphs


def is_heading(paragraph):
    """Checks if a paragraph is a heading/headline based on paragraph style

    :param paragraph:
    :return: Boolean True if paragraph is heading, else False
    """
    style = paragraph.style
    return style and style.name.startswith("Head")


def extract_content_from_docx(docx_file, doc_sections):
    """Extracts section wise text from a document where sections are the names of headings required

    :param docx_file: input document from which text to be extracted
    :param doc_sections: heading names of paragraphs/sections in document
    :return: list of dictionaries with section and context as keys
    """
    content_list = []
    document = Document(docx_file)
    for paragraphs in iterate_document_sections(document):
        for doc_section in doc_sections:
            if doc_section in paragraphs[0].text.lower():
                if doc_section == "title":
                    content_list.append(
                        {"section": "title", "context": "".join(paragraph.text for paragraph in paragraphs)})
                else:
                    context = "\n".join(paragraph.text for paragraph in paragraphs[1:])
                    if context != "":
                        content_list.append({"section": doc_section, "context": context})

    return content_list


def get_prompt_template_old(question, context, length):
    """Gives a general prompt template/ instructions to LLM, takes context as input for model

    :param question: prompt/query for llm
    :param context: input text related to the query
    :param length: max word count limit required in response
    :return: prompt template
    """
    qa_prompt = PromptTemplate(
        input_variables=["context"],
        partial_variables={"question": question, "answer_length": length},
        template="""Write an answer, strictly not more than {answer_length} words
    for the question below based on the provided context. 
    Do not break & miss the important facts from the context.
    Answer in an unbiased, comprehensive, and scholarly tone.\n\n
    {context}\n
    Question: {question}\n
    Answer: """
    )
    return qa_prompt


def get_prompt_template_faiss(question, length):
    """Gives a general prompt template/instructions to LLM, takes vectors (from faiss) as model input

    :param question: prompt/query for llm
    :param length: max word count limit required in response
    :return: prompt template
    """
    qa_prompt = PromptTemplate(
        input_variables=["question", "length"],
        # partial_variables={"question": question, "length": length},
        template="""Write a very concise answer. Provide an answer strictly not more than {length} words.
    Do not break & miss the important facts from the context.
    Avoid repetition and ensure the answer includes all the facts.
    Just provide the short answers, do not provide any explanations or opening/closing sentences.
    If no answer is found just give ''\n\n
    Question: {question}
    """
    )
    return qa_prompt.format(question=question, length=length)


def get_prompt_template(question, context, length):
    """Gives a general prompt template/ instructions to LLM, takes context as input for model

    :param question: prompt/query for llm
    :param context: input text related to the query
    :param length: max word count limit required in response
    :return: prompt template
    """
    if length != 'NotDefined':
        qa_prompt = PromptTemplate(
            input_variables=["context"],
            partial_variables={"question": question, "answer_length": length},
            template="""Write an answer, strictly not more than {answer_length} words
        for the question below based on the provided context. 
        Do not break & miss the important facts from the context.
        Answer in an unbiased, comprehensive, and scholarly tone.\n\n
        {context}\n
        Question: {question}\n
        Answer: """
        )
    else:
        qa_prompt = PromptTemplate(
            input_variables=["context"],
            partial_variables={"question": question},
            template="""Write an answer for the question below based on the provided context. 
        Do not break & miss the important facts from the context.
        Answer in an unbiased, comprehensive, and scholarly tone.\n\n
        {context}\n
        Question: {question}\n
        Answer: """
        )

    return qa_prompt


def get_prompt_template_ISD(question, length):
    """Gives a general prompt template/instructions to LLM, takes vectors (from faiss) as model input
        and does index search in documents

    :param question: prompt/query for llm
    :param length: max word count limit required in response
    :return: prompt template
    """
    if length != 'NotDefined':
        qa_prompt = PromptTemplate(
            input_variables=["question", "length"],
            template="""Write an answer, strictly not more than {length} words.
        Do not break & miss the important facts.
        Answer in an unbiased, comprehensive, and scholarly tone.\n\n
        Question: {question}\
        Answer: """
        )
        qa_prompt = qa_prompt.format(question=question, length=length)
    else:
        qa_prompt = PromptTemplate(
            input_variables=["question"],
            template="""Write an answer for the question. 
        Do not break & miss the important facts.
        Answer in an unbiased, comprehensive, and scholarly tone.\n\n
        Question: {question}\n
        Answer: """
        )
        qa_prompt = qa_prompt.format(question=question)

    return qa_prompt


def get_query_answer(context, model, max_tokens, question, length):
    """Get llm response based on prompt template

    :param context: input text for model
    :param model: llm model name
    :param max_tokens: max tokens parameter for llm
    :param question: prompt/query for llm
    :param length: max word count limit for llm response
    :return: response from llm
    """
    llm = ChatOpenAI(model_name=model, max_tokens=max_tokens, temperature=0, openai_api_key=OPENAI_API_KEY)
    prompt_template = get_prompt_template(question, context, length)
    # print(f"qa_template \n:{prompt_template}")
    chain = LLMChain(llm=llm, prompt=prompt_template)
    output = chain.run(context)
    return output


def get_length(prompt):
    """Gets max word count limit or length as defined in prompt dictionary

    :param prompt: prompt dictionary user input
    :return: dictionary with section name, question and word count
    """
    section_dict = {}
    count = 0
    while count < len(prompt.keys()):
        section_dict[prompt[count]['section'] + ' ' + prompt[count]['question']] = prompt[count]['words']
        count += 1
    return section_dict


def index_and_search_documents(file_path, query, length, model, max_tokens, embeddings):
    """Get llm response by using embeddings saved in faiss as model input

    :param file_path: input file path
    :param query: prompt/question for llm
    :param length: max word count limit for response
    :param model: llm model name
    :param max_tokens: max tokens parameter for llm
    :param embeddings: openai embeddings
    :return: llm response
    """
    if file_path.endswith(".docx"):
        pdf_file_path = file_path[:-4] + "pdf"
    else:
        pdf_file_path = file_path
    loader = PyPDFLoader(pdf_file_path)
    # loader = TextLoader(file_path,encoding='latin-1')
    pages = loader.load_and_split()
    db_faiss = FAISS.from_documents(pages, embeddings)
    prompt_template = get_prompt_template_ISD(query, length)
    print(prompt_template)
    chatbot = RetrievalQA.from_chain_type(
        llm=ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            temperature=0, model_name=model, max_tokens=max_tokens
        ),
        chain_type=chain_type,
        retriever=db_faiss.as_retriever()
    )
    response = chatbot.run(prompt_template)
    return response


def format_manuscript(text, formatting_instructions_list, OPENAI_API_KEY):
    """Restyle the text with some guidelines as mentioned in instruction list

    :param text: input text to be restyled
    :param formatting_instructions_list: list of guidelines for restyling
    :return: formatted/restyled text
    """
    openai.api_key = OPENAI_API_KEY
    formatted_chunks = []
    num_chunks = len(text) // 4096 + 1
    for i in range(num_chunks):
        start = i * 4096
        end = (i + 1) * 4096
        chunk = text[start:end]

        prompt = f'''
        You have a manuscript document that needs to be formatted according to specific guidelines.
        Text (Chunk {i + 1}/{num_chunks}): {chunk}
        Formatting instructions: {formatting_instructions_list[i]}
        Please format this part of the document according to the guideline mentioned above and apply the necessary
        changes to ensure compliance.
        '''
        response = openai.Completion.create(
            engine='text-davinci-003',
            prompt=prompt,
            max_tokens=500,
            n=1,
            stop=None,
            temperature=0,
            top_p=1.0
        )
        formatted_chunk = response.choices[0].text.strip()
        formatted_chunks.append(formatted_chunk)
    formatted_manuscript = '\n'.join(formatted_chunks)
    return formatted_manuscript


def journal_restyling(pdf_file_path, prompt, length_dict, model, max_tokens, data, instructions_list, OPENAI_API_KEY):
    """Returns dictionary of response and section/title name based on user defined prompts

    :param pdf_file_path:
    :param prompt:
    :param length_dict:
    :param model:
    :param max_tokens:
    :param data:
    :param instructions_list:
    :return:
    """
    return_dict = {}
    count = 0
    while count < len(prompt.keys()):
        title = prompt[count]['section']
        query = prompt[count]['question']
        # Check if content is present in the table
        found_in_table = False
        for i in data:
            if i['section'] == title.lower():
                context = i['context']
                length = length_dict[title + ' ' + query]
                res_text = get_query_answer(context=context, model=model, max_tokens=max_tokens, question=query,
                                            length=length)
                found_in_table = True
                break
        if not found_in_table or not res_text:
            # Perform semantic search using FAISS indexed documents
            length = length_dict[title + ' ' + query]
            res_text = index_and_search_documents(pdf_file_path, query, length, model, max_tokens, embeddings)
        count += 1
        print("Finding for Title = {} and Query = {}".format(title, query))
        if title.lower().startswith("title") is False:
            res_text = format_manuscript(res_text, instructions_list, OPENAI_API_KEY)
        if title in return_dict:
            return_dict[title].append(res_text)
        else:
            return_dict[title] = [res_text]
    return return_dict


def format_references(references_context):
    """Formats the References section in manuscript document

    :param references_context: references text from the document
    :return: list of formatted references
    """
    references = references_context.split("\n")
    new_references_list = []
    for r in references:
        if len(r.split(",")) >= 7:
            new_references_list.append(",".join(r.split(",")[0:3]) + " et al." + r.split(" ")[-1])
        else:
            new_references_list.append(r)
    return "\n".join(new_references_list)


def final_run_old(pdf_file_path, prompt, length_dict, model, max_tokens, data):
    return_dict = {}

    count = 0
    while count < len(prompt.keys()):
        title = prompt[count]['section']
        query = prompt[count]['question']

        # Check if content is present in the table
        found_in_table = False
        for i in data:

            if i['section'] == title.lower():
                context = i['context']
                length = length_dict[title + ' ' + query]
                res_text = get_query_answer(context=context, model=model, max_tokens=max_tokens, question=query,
                                            length=length)
                found_in_table = True
                break

        if not found_in_table or not res_text:
            # Perform semantic search using FAISS indexed documents
            length = length_dict[title + ' ' + query]
            res_text = index_and_search_documents(pdf_file_path, query, length, model, max_tokens, embeddings)

        count += 1
        print("Finding for Title = {} and Query = {}".format(title, query))

        if title in return_dict:
            return_dict[title].append(res_text)
        else:
            return_dict[title] = [res_text]

    return return_dict


def final_run(file_path, prompt, length_dict, model, max_tokens, data, doc_type, mapping_dict):
    """Generates response for all sections defined in prompt for a given document type

    :param file_path: source document path
    :param prompt: prompt dictionary with all the queries/prompts section wise
    :param length_dict: dictionary with word counts for sections
    :param model: openai llm model
    :param max_tokens: max tokens parameter for model
    :param data: data dictionary of source document with sections and respective contexts
    :param doc_type: document type for which content is generated
    :param mapping_dict: dictionary with a mapping of source and target section names
    :return: final response dictionary with section name and response generated
    """
    return_dict = {}

    count = 0
    print(f"Running Prompts for Document:{doc_type}\n")
    while count < len(prompt.keys()):
        title = prompt[count]['section']
        mapping_title = title
        print("prompt title:", title)
        query = prompt[count]['question']
        length = length_dict[title + ' ' + query]
        for j in range(len(mapping_dict)):
            if mapping_dict[j]['target_document_section'].lower() == title.lower():
                mapping_title = mapping_dict[j]['source_document_section']
                print("mapping title", mapping_title)

        print("Finding for Title = {} and Query = {}".format(title, query))

        # Check if content is present in the table
        found_in_table = False
        if data:
            for i in data:
                if (i['section'].lower() == title.lower()) or (i['section'].lower() == mapping_title.lower()):
                    context = i['context']
                    res_text = get_query_answer(context=context, model=model, max_tokens=max_tokens, question=query,
                                                length=length)
                    found_in_table = True
                    break

            if not found_in_table:
                # Perform semantic search using FAISS indexed documents
                print(" running index_and_search_documents from 1 for section: ", title)
                res_text = index_and_search_documents(file_path, query, length, model, max_tokens, embeddings)

        else:
            # Perform semantic search using FAISS indexed documents
            print(" running index_and_search_documents from 2")
            res_text = index_and_search_documents(file_path, query, length, model, max_tokens, embeddings)

        count += 1

        if title in return_dict:
            return_dict[title].append(res_text)
        else:
            return_dict[title] = [res_text]

    return return_dict


def calculate_cosine_similarity(source_text, target_text):
    """Calculates cosine similarity between two given embeddings

    :param source_text: text from source document
    :param target_text: text from target document
    :return: float number cosine similarity score
    """
    resp1 = openai.Embedding.create(input=[source_text], engine="text-embedding-ada-002")
    resp2 = openai.Embedding.create(input=[target_text], engine="text-embedding-ada-002")
    embedding_1 = resp1['data'][0]['embedding']
    embedding_2 = resp2['data'][0]['embedding']
    similarity_score = np.dot(embedding_1, embedding_2)
    return similarity_score


def read_pdf(filepath):
    """ function that reads a pdf file and returns its text as string

    :param filepath: input pdf file
    :return: string text format of pdf file contents
    """
    with pdfplumber.open(filepath) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text


def read_docx(filepath):
    """ function that reads a docx file and returns its text as string

    :param filepath: input pdf file
    :return: string text format of pdf file contents
    """
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]
    text = "\n".join(paragraphs)
    return text


def calculate_rouge_score(source, target):
    """ calculates rouge_1, rouge_2 and rouge_l scores for given source and target
    rouge_1 measures using matching uni-grams
    rouge_2 measures using matching bi-grams
    rouge_l measures using matching Longest Common Subsequence (LCS) words

    :param source: source document text
    :param target: target document text
    :return: rouge_1, rouge_2 and rouge_l scores
    """
    scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
    scores = scorer.score(source, target)

    rouge_1_score = scores['rouge1'].fmeasure
    rouge_2_score = scores['rouge2'].fmeasure
    rouge_l_score = scores['rougeL'].fmeasure
    return rouge_1_score, rouge_2_score, rouge_l_score


def calculate_bleu_score(source, target, weights):
    """ calculates blue_scores using matching uni-grams, bi-grams and tri-grams

    :param source: source document text
    :param target: target document text
    :param weights: parameter to change for required n-grams
    :return: bleu_1, bleu_2 and bleu_3 scores
    """
    references = [source.split()]
    hypotheses = [target.split()]

    # Calculate BLEU score
    # weights = [1.0,0,0,0]  # Use uniform weights for n-grams
    bleu_score = corpus_bleu(references, hypotheses, weights=weights)
    return bleu_score


