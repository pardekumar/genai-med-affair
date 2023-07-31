from utils import *
from config import *
from langchain.embeddings import OpenAIEmbeddings
import tiktoken
import pandas as pd
from docx import Document


# declaring variables
pdf_file = "P-Reality-X Manuscript_Draft 1_17Feb22 (PAL1144).pdf"
pdf_file_path = os.path.join(dir_name + '\\data\\input\\'+pdf_file)
file = "P-Reality-X Manuscript_Draft 1_17Feb22 (PAL1144).docx"
file_path = os.path.join(dir_name+'\\data\\input\\'+file)
doc_type = "NEJM"
input_doc = Document(file_path)
tokenizer = tiktoken.encoding_for_model(tokenizer_model)
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
prompt = pd.read_excel(excel_file)
prompt.fillna('', inplace=True)
prompt = prompt[(prompt['Document Type'] == doc_type)]
prompt.reset_index(drop=True, inplace=True)
prompt['words'].fillna('NotDefined', inplace=True)
prompt_dict = prompt.to_dict('index')
document = Document()
style = document.styles['Normal']
style.paragraph_format.line_spacing = 2.0


def get_required_sections(input_document):
    """ Get required sections from input document paragraphs

    :param input_document: document to be restyled in .docx extension
    :return: list of required sections
    """
    sections = []
    for para in input_document.paragraphs:
        if is_heading(para):
            sections.append(para.text)
    # print(len(sub_sections),sub_sections)
    new_list = []
    for i in sections[::-1]:
        if i.lower() == "abstract":
            new_list = sections[sections.index(i):]
    return [x.lower().strip() for x in new_list if x not in ('', '\n')]


required_sections = get_required_sections(input_doc)
content_list = extract_content_from_docx(file_path, required_sections)
print(content_list)

prompt.fillna('', inplace=True)
prompt = prompt.to_dict('index')
length_dict = get_length(prompt)

response = journal_restyling(pdf_file_path, prompt, length_dict, model, max_tokens, content_list, instructions_list, OPENAI_API_KEY)

first_para = input_doc.paragraphs[0]
document.add_paragraph().add_run(first_para.text).bold = True
for section, paragraphs in response.items():
    if section.title().lower() == "background":
        create_page_break(document)
        document.add_paragraph().add_run("ABSTRACT").bold = True
    if "Title" not in section.title():
        document.add_paragraph().add_run(section.title().capitalize()+": ").bold = True
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

doc_components = [i.lower() for i in doc_components]
doc_components1 = [i for i in doc_components if i not in ["table", "figure"]]

for i in range(1, len(content_list)):
    if ((content_list[i])['section']).lower().startswith("references"):
        context = format_references(content_list[i]['context'])
    else:
        context = content_list[i]['context']
        context = format_manuscript(context, instructions_list, OPENAI_API_KEY)

    if ((content_list[i])['section']).lower().startswith(tuple(doc_components)):
        create_page_break(document)
    document.add_paragraph().add_run((content_list[i]['section']).upper()).bold = True
    document.add_paragraph(context)

document.save(restyled_manuscript)
