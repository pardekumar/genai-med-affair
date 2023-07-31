import tiktoken
import pandas as pd
from config import *
from utils import *
from docx import Document

tokenizer = tiktoken.encoding_for_model("gpt-3.5-turbo")
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)


def get_table_sections_from_outline(file_path):
    """

    :param file_path: path for outline-design document
    :return: data dictionary
    """

    print(file_path)
    word_doc = Document(file_path)

    # fetch sections from the table
    table = word_doc.tables[0]
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    if data:
        df = pd.DataFrame(data)
    else:
        df = pd.DataFrame()
    #df = df.shift(1)
    df.loc[0, df.columns[0]] = df.columns[0]
    df.loc[0, df.columns[1]] = df.columns[1]
    df.rename(columns={df.columns[0]: 'section', df.columns[1]: 'context'}, inplace=True)
    df['section'] = df['section'].apply(lambda x: x.lower())
    df['context'] = df['context'].apply(lambda x: x.strip(' '))
    datadict = df.to_dict('records')
    return datadict


def doc_generation(data):

    data_dict = None
    if doc_type.lower() in ['abstract', 'manuscript', 'abstract pls']:
        print(file_path)
        data_dict = get_table_sections_from_outline(file_path)
    else:
        subsections = ["title", "abstract", "background", "methods", "results", "discussion", "conclusions",
                       "acknowledgements", "references"]
        content_list = extract_content_from_docx(file_path, subsections)
    prompt = pd.read_excel(excel_file)
    prompt.fillna('', inplace=True)
    if section_list == 'All':
        prompt = prompt[(prompt['Document Type'] == doc_type)]
    else:
        prompt = prompt[(prompt['Document Type'] == doc_type) & (prompt['section'] == section_list)]
    prompt.reset_index(drop=True, inplace=True)
    prompt['words'].fillna('NotDefined', inplace=True)
    prompt_dict = prompt.to_dict('index')
    length_dict = get_length(prompt_dict)

    mapping = pd.read_excel(mapping_excel)
    mapping = mapping[(mapping['document_type'] == doc_type)]
    mapping.reset_index(drop=True, inplace=True)
    mapping_dict = mapping.to_dict('index')

    if data_dict:
        print(" final run 1")
        response = final_run(file_path, prompt_dict, length_dict, model, max_tokens, data_dict, doc_type, mapping_dict)
    else:
        print("final run 2")
        response = final_run(file_path, prompt_dict, length_dict, model, max_tokens, content_list, doc_type,
                             mapping_dict)
    # Social Media Post

    document = Document()
    for query, paragraphs in response.items():

        document.add_heading(query.title(), level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    document.save(generated_document_path)

    return response
